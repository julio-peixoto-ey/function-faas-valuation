import base64
import os
import time
import tempfile
import logging
import re
from typing import List, Dict, Any, Optional
import fitz
import instructor
from openai import AzureOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from .utils import TokenCounter
from .model import (
    DocumentExtractionResult,
    ContractEntity,
    SeriesExtractedEntity,
    ContractEntitiesResponse,
    BulkFileUploadResponse,
    FileUploadResponse,
    DocumentModel,
)
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
import azure.functions as func
from .ocr_service import OCRService
from PIL import Image
import io
import sys

# Fix OpenMP conflict - DEVE SER ANTES de qualquer import que use OpenMP
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'
os.environ['OMP_NUM_THREADS'] = '1'

try:
    AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
    AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")
    AZURE_DEPLOYMENT_NAME = os.getenv("AZURE_DEPLOYMENT_NAME")
    AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-06-01")
    AZURE_EMBEDDING_DEPLOYMENT = os.getenv("AZURE_EMBEDDING_DEPLOYMENT")

    logging.info("Todas as variáveis de ambiente carregadas com sucesso")
except ValueError as e:
    logging.error(f"Erro na configuração: {e}")
    raise

# Configurar encoding para evitar problemas com caracteres especiais
if os.name == 'nt':  # Windows
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# Suprimir logs verbosos do EasyOCR
os.environ['EASYOCR_MODULE_PATH'] = os.getcwd()
logging.getLogger('easyocr').setLevel(logging.WARNING)

class DocumentEntityExtractor:
    """Extrator de entidades específicas de documentos usando embeddings"""

    def __init__(self):
        self.token_counter = TokenCounter()

        base_client = AzureOpenAI(
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
            api_key=AZURE_OPENAI_KEY,
            azure_deployment=AZURE_DEPLOYMENT_NAME,
            api_version=AZURE_OPENAI_API_VERSION,
        )

        self.llm = instructor.from_openai(base_client)

        self.embeddings = AzureOpenAIEmbeddings(
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
            api_key=AZURE_OPENAI_KEY,
            azure_deployment=AZURE_EMBEDDING_DEPLOYMENT,
            api_version=AZURE_OPENAI_API_VERSION,
        )

    def create_vector_store_from_chunks(
        self, document_chunks: List[Dict[str, Any]]
    ) -> FAISS:
        """Cria um banco de dados vetorial FAISS a partir dos chunks já processados"""
        logging.info(
            f"Criando banco vetorial com {len(document_chunks)} chunks pré-processados"
        )

        documents = []
        chunk_texts = []

        for chunk in document_chunks:
            chunk_content = base64.b64decode(chunk["content"]).decode("utf-8")
            chunk_texts.append(chunk_content)

            doc = Document(
                page_content=chunk_content,
                metadata={
                    "chunk_id": int(chunk["chunk_id"]),
                    "page": int(chunk["page"]),
                    "tokens": int(chunk["tokens"]),
                },
            )
            documents.append(doc)

        self.token_counter.log_embedding_usage(chunk_texts)

        vector_store = FAISS.from_documents(documents, self.embeddings)

        logging.info(f"Banco vetorial criado com {len(documents)} chunks")
        return vector_store

    def extract_all_entities(
        self, document_chunks: List[Dict[str, Any]], filename: str
    ) -> DocumentExtractionResult:
        """Extrai entidades usando Instructor para garantir JSON válido"""
        start_time = time.time()

        try:
            vector_store = self.create_vector_store_from_chunks(document_chunks)

            query = """
            Contratos financeiros: indexação de juros, spread, taxas, datas de emissão e vencimento, 
            valores nominais, cronogramas de pagamento, atualização monetária IPCA IGPM SELIC, 
            bases de cálculo 252 365 dias, fluxos de amortização, DI CDI pré-fixado, múltiplas séries.
            """

            docs = vector_store.similarity_search_with_score(query, k=40)

            if not docs:
                return self._create_empty_result(filename, start_time)

            context_parts = []
            for doc, score in docs:
                context_parts.append(
                    f"Chunk {doc.metadata.get('chunk_id', 'N/A')} (página {doc.metadata.get('page', 'N/A')}):\n{doc.page_content}"
                )

            full_context = "\n\n---\n\n".join(context_parts)
            prompt = self._build_extraction_prompt(full_context)

            response = self.llm.chat.completions.create(
                model=AZURE_DEPLOYMENT_NAME,
                response_model=ContractEntitiesResponse,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=4096,
                max_retries=3,
            )

            contract_entities = self._process_extraction_response(response, full_context)
            processing_time = int((time.time() - start_time) * 1000)
            
            entities_count = sum(1 for field in contract_entities.__dataclass_fields__ 
                               if getattr(contract_entities, field) is not None)
            logging.info(f"Extraídas {entities_count}/9 entidades para {filename}")

            return DocumentExtractionResult(
                filename=filename,
                success=True,
                contract_entities=contract_entities,
                processing_time_ms=processing_time,
                token_summary=self.token_counter.get_summary(),
            )

        except Exception as e:
            logging.error(f"Erro na extração: {e}")
            return self._create_empty_result(filename, start_time)

    def _build_extraction_prompt(self, context: str) -> str:
        return f"""
            <role>
            Você é um(a) **analista sênior de contratos financeiros**.  
            Sua tarefa é **LER** o texto abaixo e **DEVOLVER** exatamente **um** objeto
            JSON com os nove campos pedidos, **com listas** para cada campo (cada posição da lista representa uma série).
            </role>
            
            <series>
            O contexto é um documento de contrato financeiro.
            O documento pode conter uma ou mais séries.
            Cada série representa um contrato diferente.
            Cada campo deve ser uma LISTA onde cada posição corresponde a uma série.
            Se houver apenas uma série, retorne listas com um único elemento.
            Se não houver informação para uma série específica, use "NÃO ENCONTRADO" para essa posição.
            
            <how_identify_series>
            no inicio do documento, deve conter frases com séries como:
            "DAS 1ª, 2ª E 3ª SÉRIES"
            "DAS 1ª (PRIMEIRA), 2ª (SEGUNDA) e 3ª (TERCEIRA) SÉRIES,"
            
            A quantidade de séries vai definir o tamanho das listas.
            
            </how_identify_series>
            
            </series>
            
            <context>
            {context}
            </context>
            
            <rules>
            REGRAS OBRIGATÓRIAS
            1. Copie o conteúdo **exatamente como está no contrato** – não traduza
            nem reescreva números, índices ou datas.
            2. Se o item não existir para uma série, responda **"NÃO ENCONTRADO"**.
            3. Se houver mais de um valor para o mesmo item na mesma série, una-os em **uma única
            string separada por vírgulas**, mantendo a ordem em que aparecem.
            4. Retorne listas para cada campo, onde cada posição corresponde a uma série.
            5. Retorne apenas o JSON válido (sem comentários, sem texto antes ou depois).
            </rules>
            
            <items>
            ITENS QUE DEVEM SER EXTRAÍDOS COMO LISTAS:

            1. **ATUALIZAÇÃO MONETÁRIA** – lista de índices que corrigem o **principal** por série
            2. **JUROS REMUNERATÓRIOS** – lista de indexadores **principal** por série (DI, CDI, IPCA, etc.)
            3. **SPREAD FIXO** – lista de percentuais adicionais por série
            4. **BASE DE CÁLCULO** – lista de metodologias por série (252, 365, ACT/360)
            5. **DATA EMISSÃO** – lista de datas de emissão por série
            6. **DATA VENCIMENTO** – lista de datas de vencimento por série
            7. **VALOR NOMINAL UNITÁRIO** – lista de valores de face por série
            8. **FLUXOS DE PAGAMENTO** – lista de cronogramas de pagamento por série
            9. **FLUXOS PERCENTUAIS** – lista de percentuais de amortização por série
            </items>
        """

    def _process_extraction_response(self, response, context: str) -> ContractEntity:
        contract_entities = ContractEntity()
        
        for field_name, value_list in response.model_dump().items():
            if value_list and isinstance(value_list, list):
                valid_values = []
                confidences = []
                contexts = []
                
                for value in value_list:
                    if value and value.strip() and value.upper() != "NÃO ENCONTRADO":
                        valid_values.append(value.strip())
                        confidence = self._calculate_confidence(value, context)
                        confidences.append(confidence)
                        contexts.append(context[:500] + "..." if len(context) > 500 else context)
                    else:
                        valid_values.append("NÃO ENCONTRADO")
                        confidences.append(0.0)
                        contexts.append(None)
                
                if valid_values:
                    series_entity = SeriesExtractedEntity(
                        entity_type=field_name,
                        values=valid_values,
                        confidences=confidences,
                        contexts=contexts,
                    )
                    setattr(contract_entities, field_name, series_entity)
        
        return contract_entities


    def _calculate_confidence(self, value: str, context: str) -> float:
        if not value or not context:
            return 0.5
        
        if value.lower() in context.lower():
            return 0.9
        
        keywords = value.lower().split()
        found_keywords = sum(1 for keyword in keywords if keyword in context.lower())
        
        if found_keywords > 0:
            confidence = min(0.8, 0.5 + (found_keywords / len(keywords)) * 0.3)
            return round(confidence, 3)
        
        return 0.5

    def _create_empty_result(
        self, filename: str, start_time: float
    ) -> DocumentExtractionResult:
        """Cria um resultado vazio quando nenhum chunk relevante é encontrado"""
        processing_time = int((time.time() - start_time) * 1000)

        return DocumentExtractionResult(
            filename=filename,
            success=False,
            contract_entities=ContractEntity(),
            processing_time_ms=processing_time,
            token_summary=self.token_counter.get_summary(),
            error_message="Nenhum chunk relevante encontrado no documento",
        )

    def create_tabelas_faas_por_series(
        self, contract_entities: ContractEntity, filename: str
    ) -> List[List[Dict[str, Any]]]:
        """
        Cria tabelas FAAS separadas para cada série encontrada no documento.
        Retorna uma lista de tabelas, onde cada tabela corresponde a uma série.
        """
        tabelas_por_series = []
        
        try:
            # Determinar o número de séries
            num_series = self._get_series_count(contract_entities)
            logging.info(f"DEBUG: Criando tabelas FAAS para {num_series} série(s) em {filename}")
            
            if num_series == 0:
                logging.warning(f"Nenhuma série encontrada em {filename}")
                return []
            
            # Criar uma tabela para cada série
            for serie_index in range(num_series):
                logging.info(f"DEBUG: Processando série {serie_index + 1}/{num_series}")
                
                tabela_serie = self._create_tabela_for_single_series(
                    contract_entities, filename, serie_index
                )
                
                if tabela_serie:
                    tabelas_por_series.append(tabela_serie)
                    logging.info(f"Tabela criada para série {serie_index + 1} com {len(tabela_serie)} linhas")
                else:
                    logging.warning(f"Não foi possível criar tabela para série {serie_index + 1}")
            
            logging.info(f"Total de {len(tabelas_por_series)} tabelas FAAS criadas para {filename}")
            return tabelas_por_series
            
        except Exception as e:
            logging.error(f"Erro ao criar tabelas FAAS por séries para {filename}: {e}")
            import traceback
            logging.error(f"Traceback completo: {traceback.format_exc()}")
            return []

    def create_rows_faas_resumo_por_series(
        self, contract_entities: ContractEntity, filename: str
    ) -> List[Dict[str, Any]]:
        """
        Cria rows de resumo FAAS separadas para cada série encontrada no documento.
        Retorna uma lista de rows, onde cada row corresponde a uma série.
        """
        rows_resumo = []
        
        try:
            # Determinar o número de séries
            num_series = self._get_series_count(contract_entities)
            logging.info(f"DEBUG: Criando rows de resumo para {num_series} série(s) em {filename}")
            
            if num_series == 0:
                logging.warning(f"Nenhuma série encontrada em {filename}")
                return []
            
            # Criar uma row para cada série
            for serie_index in range(num_series):
                logging.info(f"DEBUG: Processando resumo série {serie_index + 1}/{num_series}")
                
                row_resumo = self._create_resumo_for_single_series(
                    contract_entities, filename, serie_index
                )
                
                if row_resumo:
                    rows_resumo.append(row_resumo)
                    logging.info(f"Row de resumo criada para série {serie_index + 1}")
                else:
                    logging.warning(f"Não foi possível criar row de resumo para série {serie_index + 1}")
            
            logging.info(f"Total de {len(rows_resumo)} rows de resumo criadas para {filename}")
            return rows_resumo
            
        except Exception as e:
            logging.error(f"Erro ao criar rows de resumo por séries para {filename}: {e}")
            import traceback
            logging.error(f"Traceback completo: {traceback.format_exc()}")
            return []

    def _get_series_count(self, contract_entities: ContractEntity) -> int:
        """Determina o número de séries baseado nas entidades extraídas"""
        max_series = 0
        
        # Verificar cada entidade e pegar o maior número de séries
        for field_name in contract_entities.__dataclass_fields__:
            entity = getattr(contract_entities, field_name)
            if entity and hasattr(entity, 'series_count'):
                max_series = max(max_series, entity.series_count)
        
        return max_series

    def _create_tabela_for_single_series(
        self, contract_entities: ContractEntity, filename: str, serie_index: int
    ) -> List[Dict[str, Any]]:
        """Cria uma tabela FAAS para uma série específica"""
        tabela_faas = []
        
        try:
            # Extrair valores específicos da série
            fluxos_pagamento = self._extract_dates_from_series(
                contract_entities.fluxos_pagamento, serie_index
            )
            fluxos_percentuais = self._extract_percentages_from_series(
                contract_entities.fluxos_percentuais, serie_index
            )
            
            valor_nominal = self._get_value_for_series(
                contract_entities.valor_nominal_unitario, serie_index, ""
            )
            index_info = self._get_value_for_series(
                contract_entities.juros_remuneratorios, serie_index, ""
            )
            
            serie_id = f"{filename}_serie_{serie_index + 1}"
            
            if not fluxos_pagamento:
                logging.info(f"DEBUG: Nenhum fluxo de pagamento para série {serie_index + 1}, criando linha básica")
                
                data_inicio = self._get_date_for_series(
                    contract_entities.data_emissao, serie_index, True
                )
                data_vencimento = self._get_date_for_series(
                    contract_entities.data_vencimento, serie_index, False
                )
                
                linha_faas = {
                    "Código": serie_id,
                    "Início": data_inicio,
                    "Vencimento": data_vencimento,
                    "Valor Nominal": valor_nominal,
                    "Valor Atualizado": "",
                    "% Amort": "100,00%",
                    "Amort. Nominal": "",
                    "Amort. Atual.": "",
                    "Amort. extra.": "",
                    "Remuneração": index_info,
                    "D": "",
                }
                tabela_faas.append(linha_faas)
                logging.info(f"Linha básica criada para série {serie_index + 1}")
            else:
                logging.info(f"DEBUG: Criando {len(fluxos_pagamento)} linhas para série {serie_index + 1}")
                
                for i, data_pagamento in enumerate(fluxos_pagamento):
                    percentual_amortizacao = (
                        fluxos_percentuais[i] if i < len(fluxos_percentuais) else ""
                    )
                    
                    if i == 0:
                        data_inicio = self._get_date_for_series(
                            contract_entities.data_emissao, serie_index, True
                        )
                    else:
                        data_inicio = fluxos_pagamento[i - 1]
                    
                    linha_faas = {
                        "Código": serie_id,
                        "Início": data_inicio,
                        "Vencimento": data_pagamento,
                        "Valor Nominal": valor_nominal,
                        "Valor Atualizado": "",
                        "% Amort": percentual_amortizacao,
                        "Amort. Nominal": "",
                        "Amort. Atual.": "",
                        "Amort. extra.": "",
                        "Remuneração": index_info,
                        "D": "",
                    }
                    
                    tabela_faas.append(linha_faas)
            
            return tabela_faas
            
        except Exception as e:
            logging.error(f"Erro ao criar tabela para série {serie_index + 1}: {e}")
            return []

    def _create_resumo_for_single_series(
        self, contract_entities: ContractEntity, filename: str, serie_index: int
    ) -> Dict[str, Any]:
        """Cria uma row de resumo para uma série específica"""
        try:
            # Extrair valores específicos da série
            data_inicio = self._get_date_for_series(
                contract_entities.data_emissao, serie_index, True
            )
            data_vencimento = self._get_date_for_series(
                contract_entities.data_vencimento, serie_index, False
            )
            
            index_info = self._get_value_for_series(
                contract_entities.juros_remuneratorios, serie_index, ""
            )
            
            serie_id = f"{filename}_serie_{serie_index + 1}"
            
            linha_resumo = {
                "Nome do Arquivo": serie_id,
                "Fundo": "Agente",
                "Link A": "",
                "Index": index_info,
                "Aplicação": "",
                "Emissão": data_inicio,
                "Vencimento": data_vencimento,
                "Quantidade": "",
                "PU Mercado": "",
                "PU Custo": "",
                "Saldo": "",
            }
            
            return linha_resumo
            
        except Exception as e:
            logging.error(f"Erro ao criar resumo para série {serie_index + 1}: {e}")
            return {}

    def _get_value_for_series(
        self, entity: Optional[SeriesExtractedEntity], serie_index: int, default: str = ""
    ) -> str:
        """Extrai valor de uma entidade para uma série específica"""
        if not entity or not entity.values or serie_index >= len(entity.values):
            return default
        
        value = entity.values[serie_index]
        return value if value and value.upper() != "NÃO ENCONTRADO" else default

    def _get_date_for_series(
        self, data_emissao_entity: Optional[SeriesExtractedEntity], serie_index: int, is_first: bool
    ) -> str:
        """Extrai a data de emissão para uma série específica"""
        value = self._get_value_for_series(data_emissao_entity, serie_index, "")
        if not value:
            return ""
        
        # Se há vírgulas, pegar a primeira data
        dates = value.split(",")
        first_date = dates[0].strip() if is_first else dates[-1].strip()
        
        return self._normalize_date_format(first_date)

    def _extract_dates_from_series(
        self, fluxos_entity: Optional[SeriesExtractedEntity], serie_index: int
    ) -> List[str]:
        """Extrai lista de datas dos fluxos de pagamento para uma série específica"""
        value = self._get_value_for_series(fluxos_entity, serie_index, "")
        if not value:
            return []
        
        date_pattern = r"\d{1,2}/\d{1,2}/\d{2,4}"
        dates = re.findall(date_pattern, value)
        
        if not dates:
            # Tentativa flexível
            parts = value.split(",")
            for part in parts:
                part = part.strip()
                flexible_pattern = r"\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}"
                found_dates = re.findall(flexible_pattern, part)
                if found_dates:
                    dates.extend(found_dates)
        
        return dates

    def _extract_percentages_from_series(
        self, fluxos_entity: Optional[SeriesExtractedEntity], serie_index: int
    ) -> List[str]:
        """Extrai lista de percentuais dos fluxos percentuais para uma série específica"""
        value = self._get_value_for_series(fluxos_entity, serie_index, "")
        if not value:
            return []
        
        percentages = [p.strip() for p in value.split(",")]
        # Corrigir formato dos percentuais
        formatted_percentages = []
        for perc in percentages:
            if perc and perc != "NÃO ENCONTRADO":
                # Garantir formato correto (X,XXXX%)
                if '%' in perc:
                    # Se já tem %, manter
                    formatted_percentages.append(perc)
                else:
                    # Se não tem %, adicionar
                    formatted_percentages.append(f"{perc}%")
            else:
                formatted_percentages.append(perc)
        
        return formatted_percentages

    def _normalize_date_format(self, date_str: str) -> str:
        if not date_str or re.match(r"\d{1,2}/\d{1,2}/\d{4}", date_str):
            return date_str
        
        months = {
            "janeiro": "01", "fevereiro": "02", "março": "03", "abril": "04",
            "maio": "05", "junho": "06", "julho": "07", "agosto": "08",
            "setembro": "09", "outubro": "10", "novembro": "11", "dezembro": "12",
        }
        
        for month_name, month_num in months.items():
            if month_name in date_str.lower():
                parts = date_str.split()
                if len(parts) >= 3:
                    day = parts[0].strip()
                    year = parts[-1].strip()
                    return f"{day.zfill(2)}/{month_num}/{year}"
        
        return date_str


class DocumentTextExtractorService:
    """Serviço responsável pela extração de texto de arquivos PDF, TXT, DOC e DOCX"""

    def __init__(self, req: func.HttpRequest):
        self.supported_extensions = [".pdf", ".txt", ".doc", ".docx"]
        self.token_counter = TokenCounter()
        self.req = req
        self.splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            encoding_name="o200k_base",
            chunk_size=800,
            chunk_overlap=200,
            separators=["\n\n", "\n"],
        )
        
        # Inicializar serviço OCR
        self.ocr_service = OCRService()

    def extract_text_from_pdf_with_ocr(self, file_content: bytes, filename: str) -> List[DocumentModel]:
        """Extrai texto de PDF usando OCR (delegando para OCRService)"""
        try:
            return self.ocr_service.extract_text_with_ocr(file_content, filename)
        except Exception as e:
            logging.error(f"Erro na extração OCR para {filename}: {str(e)}")
            logging.info(f"Tentando extração normal como fallback para {filename}")
            return self.extract_text_from_pdf_normal(file_content, filename)

    def extract_text_from_pdf_normal(self, file_content: bytes, filename: str) -> List[DocumentModel]:
        """Extração normal de PDF (texto embutido)"""
        documents = []
        temp_path = None

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name

            doc = fitz.open(temp_path)

            logging.info(f"Processando PDF com extração normal: {filename} com {len(doc)} páginas")

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()

                if text.strip():
                    document = DocumentModel(
                        page_content=text.strip(),
                        page_number=page_num + 1,
                        source=filename,
                        file_type="pdf",
                        character_count=len(text.strip()),
                        token_count=self._estimate_tokens(text.strip()),
                    )
                    documents.append(document)

            doc.close()
            logging.info(f"Extração normal concluída: {len(documents)} páginas com texto de {filename}")

        except Exception as e:
            logging.error(f"Erro ao extrair texto do PDF {filename}: {str(e)}")
            raise Exception(f"Falha na extração de texto: {str(e)}")

        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception as e:
                    logging.warning(f"Erro ao remover arquivo temporário: {str(e)}")

        return documents

    def extract_text_from_pdf(self, file_content: bytes, filename: str) -> List[DocumentModel]:
        """Extrai texto de arquivos PDF, detectando automaticamente se precisa de OCR"""
        
        # Verificar se o PDF contém texto extraível
        is_text_extractable = self.ocr_service.is_pdf_text_extractable(file_content, filename)
        
        if not is_text_extractable:
            logging.info(f"{filename}: PDF detectado como escaneado, usando OCR")
            return self.extract_text_from_pdf_with_ocr(file_content, filename)
        
        try:
            documents = self.extract_text_from_pdf_normal(file_content, filename)
            
            if not documents:
                logging.info(f"{filename}: Extração normal não retornou texto, tentando OCR")
                return self.extract_text_from_pdf_with_ocr(file_content, filename)
            
            total_chars = sum(len(doc.page_content) for doc in documents)
            avg_chars_per_page = total_chars / len(documents) if documents else 0
            
            if avg_chars_per_page < 50:
                logging.info(f"{filename}: Texto extraído muito escasso ({avg_chars_per_page:.1f} chars/página), tentando OCR")
                try:
                    ocr_documents = self.extract_text_from_pdf_with_ocr(file_content, filename)
                    
                    if ocr_documents:
                        ocr_total_chars = sum(len(doc.page_content) for doc in ocr_documents)
                        if ocr_total_chars > total_chars * 1.5:
                            logging.info(f"{filename}: OCR produziu mais texto ({ocr_total_chars} vs {total_chars}), usando OCR")
                            return ocr_documents
                except Exception as e:
                    logging.warning(f"OCR falhou, usando extração normal: {str(e)}")
            
            return documents
            
        except Exception as e:
            logging.error(f"Erro na extração normal de {filename}: {str(e)}")
            logging.info(f"Tentando OCR como fallback para {filename}")
            return self.extract_text_from_pdf_with_ocr(file_content, filename)

    def process_file_upload(self) -> BulkFileUploadResponse:
        """Processa upload de múltiplos arquivos com otimizações para documentos grandes"""
        files_data = self._extract_files_from_request()

        if not files_data:
            raise ValueError("Nenhum arquivo foi enviado")

        processed_files = []
        total_size_mb = sum(len(content) for content, _ in files_data) / (1024 * 1024)
        
        logging.info(f"Iniciando processamento: {len(files_data)} arquivo(s), {total_size_mb:.1f}MB total")
        
        if len(files_data) > 10:
            logging.warning(f"Muitos arquivos ({len(files_data)}), pode haver timeout")
        
        if total_size_mb > 200:
            logging.warning(f"Volume muito grande ({total_size_mb:.1f}MB), pode haver problemas de memória")

        for idx, (file_content, filename) in enumerate(files_data):
            file_size_mb = len(file_content) / (1024 * 1024)
            
            if not self.is_supported_file(filename):
                supported_formats = ", ".join(self.supported_extensions)
                raise ValueError(
                    f"Arquivo {filename}: Apenas arquivos {supported_formats} são suportados"
                )

            logging.info(f"Arquivo [{idx+1}/{len(files_data)}] {filename} ({file_size_mb:.1f}MB)")

            try:
                file_response = self._process_single_file(file_content, filename)
                processed_files.append(file_response)

            except Exception as e:
                logging.error(f"Erro ao processar {filename}: {str(e)}")
                error_response = FileUploadResponse(
                    success=False,
                    filename=filename,
                    file_type=self._get_file_extension(filename).replace(".", ""),
                    documents_count=0,
                    total_characters=0,
                    total_tokens=0,
                    processing_time_ms=0,
                    token_summary=self.token_counter.get_summary(),
                    documents=[],
                    message=f"Erro ao processar arquivo: {str(e)}",
                )
                processed_files.append(error_response)

        success = any(file_resp.success for file_resp in processed_files)
        
        successful_files = sum(1 for f in processed_files if f.success)
        logging.info(f"Processamento finalizado: {successful_files}/{len(files_data)} arquivos processados com sucesso")
        
        response = BulkFileUploadResponse(success=success, files=processed_files)
        return response

    def _process_single_file(
        self, file_content: bytes, filename: str
    ) -> FileUploadResponse:
        file_extension = self._get_file_extension(filename)
        
        if file_extension == ".pdf":
            documents = self.extract_text_from_pdf(file_content, filename)
        elif file_extension == ".txt":
            documents = self.extract_text_from_txt(file_content, filename)
        elif file_extension in [".doc", ".docx"]:
            documents = self.extract_text_from_docx(file_content, filename)
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {file_extension}")
            
        all_texts = [doc.page_content for doc in documents]
        all_chunks = self.split_pages(all_texts)
        total_tokens = self.token_counter.log_embedding_usage(all_chunks)
        token_summary = self.token_counter.get_summary()

        chunk_data = []
        for idx, chunk in enumerate(all_chunks):
            chunk_data.append(
                {
                    "chunk_id": idx + 1,
                    "content": base64.b64encode(chunk.encode("utf-8")).decode("utf-8"),
                    "tokens": len(chunk.split()),
                    "page": self._get_chunk_page(idx, len(documents)),
                }
            )

        response = FileUploadResponse(
            success=True,
            filename=filename,
            file_type=file_extension.replace(".", ""),
            documents_count=len(chunk_data),
            total_characters=sum(len(chunk) for chunk in all_chunks),
            total_tokens=total_tokens,
            processing_time_ms=0,
            token_summary=token_summary,
            documents=chunk_data,
        )

        return response

    def _get_file_extension(self, filename: str) -> str:
        """Extrai a extensão do arquivo"""
        return os.path.splitext(filename.lower())[1]

    def extract_text_from_txt(
        self, file_content: bytes, filename: str
    ) -> List[DocumentModel]:
        """Extrai texto de arquivos TXT"""
        documents = []
        
        try:
            text = None
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise Exception("Não foi possível decodificar o arquivo TXT")

            logging.info(f"Processando TXT: {filename}")

            if text.strip():
                document = DocumentModel(
                    page_content=text.strip(),
                    page_number=1,
                    source=filename,
                    file_type="txt",
                    character_count=len(text.strip()),
                    token_count=self._estimate_tokens(text.strip()),
                )
                documents.append(document)

            logging.info(f"Extração TXT concluída: {filename}")

        except Exception as e:
            logging.error(f"Erro ao extrair texto do TXT {filename}: {str(e)}")
            raise Exception(f"Falha na extração de texto: {str(e)}")

        if not documents:
            raise Exception("Nenhum texto foi extraído do arquivo TXT")

        return documents

    def extract_text_from_docx(
        self, file_content: bytes, filename: str
    ) -> List[DocumentModel]:
        """Extrai texto de arquivos DOC e DOCX"""
        documents = []

        try:
            from docx import Document as DocxDocument
            import io

            if filename.lower().endswith('.docx'):
                doc_stream = io.BytesIO(file_content)
                doc = DocxDocument(doc_stream)
                
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
                
                text = '\n'.join(full_text)
                
            else:
                try:
                    text = file_content.decode('utf-8', errors='ignore')
                except Exception as e:
                    logging.error(f"Erro ao decodificar DOCX/DOC {filename}: {str(e)}")
                    text = file_content.decode('latin-1', errors='ignore')

            logging.info(f"Processando DOCX/DOC: {filename}")

            if text.strip():
                document = DocumentModel(
                    page_content=text.strip(),
                    page_number=1,
                    source=filename,
                    file_type="docx" if filename.lower().endswith('.docx') else "doc",
                    character_count=len(text.strip()),
                    token_count=self._estimate_tokens(text.strip()),
                )
                documents.append(document)

            logging.info(f"Extração DOCX/DOC concluída: {filename}")

        except Exception as e:
            logging.error(f"Erro ao extrair texto do DOCX/DOC {filename}: {str(e)}")
            raise Exception(f"Falha na extração de texto: {str(e)}")

        if not documents:
            raise Exception("Nenhum texto foi extraído do arquivo DOCX/DOC")

        return documents

    def _get_chunk_page(self, chunk_index: int, total_pages: int) -> int:
        if total_pages == 0:
            return 1
        page = (chunk_index % total_pages) + 1
        return page

    def _estimate_tokens(self, text: str) -> int:
        return len(text) // 4

    def is_supported_file(self, filename: str) -> bool:
        file_extension = self._get_file_extension(filename)
        return file_extension in self.supported_extensions

    def _extract_files_from_request(self) -> List[tuple[bytes, str]]:
        files_data = []

        files = self.req.files
        if files:
            for key, file_item in files.items():
                content = file_item.read()
                filename = file_item.filename or f"arquivo_{key}.pdf"
                files_data.append((content, filename))
            return files_data

        try:
            body_json = self.req.get_json()
            if body_json:
                if "files" in body_json and isinstance(body_json["files"], list):
                    for file_info in body_json["files"]:
                        if "file_content" in file_info:
                            file_content = base64.b64decode(file_info["file_content"])
                            filename = file_info.get("filename", "document.pdf")
                            files_data.append((file_content, filename))
                    return files_data

                elif "file_content" in body_json:
                    file_content = base64.b64decode(body_json["file_content"])
                    filename = body_json.get("filename", "document.pdf")
                    files_data.append((file_content, filename))
                    return files_data

        except Exception as e:
            logging.error(f"Erro ao processar JSON: {str(e)}")

        return files_data

    def split_pages(self, pages: list[str]) -> list[str]:
        joined = "\n".join(pages)
        return self.splitter.split_text(joined)
